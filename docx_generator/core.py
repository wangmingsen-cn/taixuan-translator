"""
太玄智译 - Word文档生成核心
基于 python-docx，符合 GB/T 7713.1-2006 出版标准
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from loguru import logger

from taixuan_translator.core.config import BilingualMode, get_settings
from taixuan_translator.core.exceptions import DocxGenerationError
from taixuan_translator.core.utils import timeit
from taixuan_translator.pdf_parser.core_parser import DocumentContent, TextBlock


class DocxGenerator:
    """
    Word文档生成器
    
    支持三种输出模式：
    1. 仅译文（translated_only）
    2. 双语对照 - 段落交替（interleaved）
    3. 双语对照 - 左右分栏（side_by_side）
    4. 双语对照 - 脚注模式（footnote）
    """

    def __init__(self) -> None:
        self.settings = get_settings()
        self.docx_cfg = self.settings.docx

    @timeit
    def generate(
        self,
        document: DocumentContent,
        translations: dict[int, str],  # {block_reading_order: translated_text}
        output_path: str | Path,
        bilingual_mode: Optional[BilingualMode] = None,
        translated_only: bool = False,
    ) -> Path:
        """
        生成Word文档。
        
        Args:
            document: PDF解析结果
            translations: 翻译结果字典 {阅读顺序索引: 译文}
            output_path: 输出文件路径
            bilingual_mode: 双语模式（None=使用配置默认值）
            translated_only: 是否只输出译文
            
        Returns:
            生成的Word文件路径
        """
        try:
            from docx import Document  # type: ignore[import]
            from docx.shared import Pt, Cm, RGBColor  # type: ignore[import]
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
        except ImportError as e:
            raise DocxGenerationError("python-docx 未安装，请运行: pip install python-docx") from e

        mode = bilingual_mode or self.settings.bilingual_mode
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"开始生成Word文档: {output_path.name}, 模式={mode}")

        try:
            doc = Document()
            self._apply_document_styles(doc)
            self._set_page_margins(doc)

            # 添加文档标题
            source_name = Path(document.file_path).stem
            title_para = doc.add_heading(source_name, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[attr-defined]

            # 逐页生成内容
            for page in document.pages:
                if page.page_number > 1:
                    # 页面分隔注释（可选）
                    pass

                for block in page.text_blocks:
                    translated = translations.get(block.reading_order, "")

                    if translated_only:
                        self._add_translated_block(doc, block, translated)
                    elif mode == BilingualMode.INTERLEAVED:
                        self._add_interleaved_block(doc, block, translated)
                    elif mode == BilingualMode.FOOTNOTE:
                        self._add_footnote_block(doc, block, translated)
                    else:
                        # 默认：段落交替
                        self._add_interleaved_block(doc, block, translated)

            doc.save(str(output_path))
            logger.info(f"Word文档生成完成: {output_path}")
            return output_path

        except DocxGenerationError:
            raise
        except Exception as e:
            raise DocxGenerationError(f"Word文档生成失败: {output_path.name}", detail=str(e)) from e

    def _apply_document_styles(self, doc: object) -> None:
        """应用文档样式（GB/T 7713.1-2006）"""
        from docx.shared import Pt, RGBColor  # type: ignore[import]
        from docx.oxml.ns import qn  # type: ignore[import]

        cfg = self.docx_cfg
        styles = doc.styles  # type: ignore[attr-defined]

        # 正文样式
        normal_style = styles["Normal"]
        normal_style.font.name = cfg.font_name_en
        normal_style.font.size = Pt(cfg.font_size_body)
        # 中文字体
        normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), cfg.font_name_cn)

        # 标题样式
        for level, size in [
            ("Heading 1", cfg.font_size_h1),
            ("Heading 2", cfg.font_size_h2),
            ("Heading 3", cfg.font_size_h3),
        ]:
            try:
                h_style = styles[level]
                h_style.font.size = Pt(size)
                h_style.font.bold = True
            except KeyError:
                pass

    def _set_page_margins(self, doc: object) -> None:
        """设置页边距（GB/T标准）"""
        from docx.shared import Cm  # type: ignore[import]

        cfg = self.docx_cfg
        for section in doc.sections:  # type: ignore[attr-defined]
            section.top_margin = Cm(cfg.page_margin_top)
            section.bottom_margin = Cm(cfg.page_margin_bottom)
            section.left_margin = Cm(cfg.page_margin_left)
            section.right_margin = Cm(cfg.page_margin_right)

    def _add_interleaved_block(
        self, doc: object, block: TextBlock, translated: str
    ) -> None:
        """段落交替模式：原文段落 + 译文段落"""
        from docx.shared import RGBColor, Pt  # type: ignore[import]

        # 原文（灰色，小字）
        if block.text.strip():
            orig_para = doc.add_paragraph()  # type: ignore[attr-defined]
            orig_run = orig_para.add_run(block.text)
            orig_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            orig_run.font.size = Pt(self.docx_cfg.font_size_body * 0.9)

        # 译文（正常黑色）
        if translated.strip():
            trans_para = doc.add_paragraph()  # type: ignore[attr-defined]
            if block.block_type == "title":
                trans_para.style = "Heading 2"
            trans_run = trans_para.add_run(translated)
            trans_run.font.size = Pt(self.docx_cfg.font_size_body)

    def _add_translated_block(
        self, doc: object, block: TextBlock, translated: str
    ) -> None:
        """仅译文模式"""
        if not translated.strip():
            return
        if block.block_type == "title":
            doc.add_heading(translated, level=2)  # type: ignore[attr-defined]
        else:
            doc.add_paragraph(translated)  # type: ignore[attr-defined]

    def _add_footnote_block(
        self, doc: object, block: TextBlock, translated: str
    ) -> None:
        """脚注模式：译文正文 + 原文脚注（python-docx脚注支持有限，降级为括号注）"""
        if translated.strip():
            para = doc.add_paragraph()  # type: ignore[attr-defined]
            para.add_run(translated)
            if block.text.strip():
                note_run = para.add_run(f"  [{block.text[:80]}...]")
                from docx.shared import RGBColor, Pt  # type: ignore[import]
                note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                note_run.font.size = Pt(8)
